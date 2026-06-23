from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ARQUIVO_SAIDA = "Documento_Metodos_Ordenacao.docx"


def set_font(run, name="Calibri", size=11, bold=False, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")

    grid = tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[idx]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def style_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for style_name, size, color, before, after in [
        ("Heading 1", 16, "2E74B5", 18, 10),
        ("Heading 2", 13, "2E74B5", 14, 7),
        ("Heading 3", 12, "1F4D78", 10, 5),
    ]:
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.25


def add_title(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Metodos de Ordenacao: Shell, Heap e Merge")
    set_font(r, size=20, bold=True, color="0B2545")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(14)
    r = p.add_run("Documento de apoio com explicacao passo a passo")
    set_font(r, size=11, color="555555")


def add_summary_table(doc):
    doc.add_heading("Resumo comparativo", level=1)
    table = doc.add_table(rows=1, cols=4)
    headers = ["Metodo", "Ideia central", "Melhor caso", "Caso medio/pior caso"]
    for idx, text in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.text = text
        set_cell_shading(cell, "E8EEF5")
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                set_font(run, bold=True, color="0B2545")

    rows = [
        ["Shell", "Insertion sort com intervalos decrescentes.", "Depende dos intervalos", "Varia conforme a sequencia de gaps"],
        ["Heap", "Transforma o vetor em heap maximo e remove a raiz.", "O(n log n)", "O(n log n)"],
        ["Merge", "Divide o vetor, ordena partes menores e intercala.", "O(n log n)", "O(n log n)"],
    ]
    for row_data in rows:
        cells = table.add_row().cells
        for idx, text in enumerate(row_data):
            cells[idx].text = text

    set_table_width(table, [1440, 3600, 1800, 2520])


def add_numbered_steps(doc, steps):
    for step in steps:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.left_indent = Inches(0.375)
        p.paragraph_format.first_line_indent = Inches(-0.188)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.25
        p.add_run(step)


def add_method_shell(doc):
    doc.add_heading("Shell Sort", level=1)
    doc.add_paragraph(
        "O Shell Sort e uma melhoria do Insertion Sort. Em vez de comparar apenas elementos vizinhos, "
        "ele compara elementos separados por um intervalo chamado gap. Esse intervalo diminui ate chegar a 1."
    )
    doc.add_heading("Passo a passo", level=2)
    add_numbered_steps(doc, [
        "Escolha um intervalo inicial, normalmente metade do tamanho do vetor.",
        "Compare elementos separados por esse intervalo.",
        "Se o elemento da esquerda for maior que o da direita, troque-os.",
        "Percorra o vetor realizando uma ordenacao por insercao considerando o intervalo atual.",
        "Reduza o intervalo, geralmente dividindo por 2.",
        "Repita o processo ate o intervalo chegar a 1.",
        "Com gap igual a 1, execute a ultima passagem como um Insertion Sort comum, mas com o vetor ja parcialmente organizado.",
    ])
    doc.add_heading("Caracteristicas", level=2)
    doc.add_paragraph(
        "E um metodo interno, pois ordena no proprio vetor. Seu desempenho depende muito da sequencia de gaps escolhida."
    )


def add_method_heap(doc):
    doc.add_heading("Heap Sort", level=1)
    doc.add_paragraph(
        "O Heap Sort usa uma estrutura chamada heap. Para ordenar em ordem crescente, normalmente e usado um heap maximo, "
        "em que o maior valor fica sempre na raiz."
    )
    doc.add_heading("Passo a passo", level=2)
    add_numbered_steps(doc, [
        "Organize o vetor para que ele respeite a propriedade de heap maximo.",
        "Nesse heap, cada pai deve ser maior ou igual aos seus filhos.",
        "Troque a raiz, que contem o maior elemento, com a ultima posicao do vetor ainda nao ordenada.",
        "Considere essa ultima posicao como parte ja ordenada.",
        "Reduza o tamanho logico do heap em uma unidade.",
        "Reorganize o heap chamando o procedimento heapify a partir da raiz.",
        "Repita as trocas e reorganizacoes ate restar apenas um elemento no heap.",
    ])
    doc.add_heading("Caracteristicas", level=2)
    doc.add_paragraph(
        "Tem complexidade O(n log n) e nao precisa de vetor auxiliar grande. E eficiente para grandes quantidades de dados."
    )


def add_method_merge(doc):
    doc.add_heading("Merge Sort", level=1)
    doc.add_paragraph(
        "O Merge Sort aplica a estrategia dividir para conquistar. Ele divide o vetor em partes menores, ordena essas partes "
        "e depois junta os resultados por intercalacao."
    )
    doc.add_heading("Passo a passo", level=2)
    add_numbered_steps(doc, [
        "Divida o vetor ao meio.",
        "Continue dividindo cada metade ate obter partes com um unico elemento.",
        "Considere que uma parte com um elemento ja esta ordenada.",
        "Compare os primeiros elementos de duas partes ordenadas.",
        "Copie para um vetor auxiliar o menor elemento entre as duas partes.",
        "Repita a comparacao ate todos os elementos das duas partes serem copiados em ordem.",
        "Substitua o trecho original pelo trecho intercalado.",
        "Continue juntando as partes ate que o vetor inteiro esteja ordenado.",
    ])
    doc.add_heading("Caracteristicas", level=2)
    doc.add_paragraph(
        "Possui complexidade O(n log n) em praticamente todos os casos. Sua principal desvantagem e o uso de memoria auxiliar."
    )


def add_closing(doc):
    doc.add_heading("Conclusao", level=1)
    doc.add_paragraph(
        "Shell Sort e simples e melhora o Insertion Sort ao trabalhar com intervalos. Heap Sort organiza os dados usando a "
        "propriedade de heap e mantem bom desempenho sem grande memoria extra. Merge Sort e muito previsivel em desempenho, "
        "mas precisa de espaco auxiliar para realizar as intercalacoes."
    )


def main():
    doc = Document()
    style_document(doc)
    add_title(doc)
    doc.add_paragraph(
        "Este documento descreve, em linguagem direta, como os metodos de ordenacao Shell, Heap e Merge executam suas etapas."
    )
    add_summary_table(doc)
    add_method_shell(doc)
    add_method_heap(doc)
    add_method_merge(doc)
    add_closing(doc)
    doc.save(ARQUIVO_SAIDA)


if __name__ == "__main__":
    main()
